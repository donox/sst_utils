import re
import csv
from docx import Document
from unidecode import unidecode
# from utilities.fsa import Rule_Parse_FSM
# from create_web_pages import minutes as minutes
from utilities.minutes_parser import Minutes


class ManageDocx(object):
    def __init__(self, file):
        self.file = file
        self.multiline = False          # Single line of result may require multiple physical lines read

    def _make_entry_re(self, entry):
        res = r'^(' + entry + '):?(.*$)'
        return re.compile(res)

    def _make_subentry_re(self, entry):
        res = r'^\s+(' + entry + '):?(.*$)'
        return re.compile(res)

    def do_alexa_commands(self):
        try:
            doc = Document(self.file)
            elements = []
            for para in doc.paragraphs:
                text = unidecode(para.text)
                elements.append(self._parse_alexa_commands(text))
            return elements
        except Exception as e:
            print(e)

    def _parse_alexa_commands(self, line):
        re_1 = re.compile(r'^"(.+)".*$')
        re_2 = re.compile(r'^(.+):.*"(.+)".*(\S*)$')
        res = re.match(re_1, line)
        if res:
            return ("RE1", res.groups()[0])
        else:
            res = re.match(re_2, line)
            if res:
                g = res.groups()
                if g[2]:
                    return "RE3", g[0], g[1], g[2]
                else:
                    return "RE2", g[0], g[1]
            else:
                return "RE0", line

    def set_multiline(self, value):
        self.multiline = value

    def _make_re_for_minutes(self, target_directory, org):
        with open(target_directory + org + ' Minutes Format.csv') as csvfile:
            rdr = csv.reader(csvfile)
            entries = []
            current_top = None
            for line_type, level, parse_code, layout in rdr:
                if line_type != 'Entry Name' and line_type != '':
                    if line_type == 'END':
                        break
                    elif level == '1':
                        l2 = []
                        entries.append((line_type, parse_code, layout, l2))
                    else:
                        l2.append((line_type, parse_code, layout))
            res_list = []
            for entry, pc, ly, sublist in entries:
                entry_re = self._make_entry_re(entry)
                sub_res = []
                for subentry, pc, ly in sublist:
                    sub_re = self._make_subentry_re(subentry)
                    sub_res.append((sub_re, [], pc, ly))
                res_list.append((entry_re, sub_res, pc, ly))
            return res_list

    def do_minutes(self, work_directory, org):
        def local_reader(doc):
            def do_read():
                for para in doc.paragraphs:
                    text = unidecode(para.text).rstrip()
                    yield text;
            return do_read
        try:
            doc_structure = self._make_re_for_minutes(work_directory, org)
            doc = Document(self.file)
            parser = Minutes(local_reader(doc), doc_structure, None, None)
            context = parser.run()
            return context
        except Exception as e:
            print(e)
